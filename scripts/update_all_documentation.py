import os
import json
import yaml
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

BRAND_DARK = RGBColor(14, 54, 70)      # #0E3646
BRAND_PRIMARY = RGBColor(245, 124, 0)   # #F57C00
TEXT_MUTED = RGBColor(100, 116, 139)    # #64748B
BG_LIGHT = "F8FAFC"
BORDER_COLOR = "CBD5E1"

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = BRAND_DARK
            run.font.size = Pt(20)
            run.bold = True
        elif level == 2:
            run.font.color.rgb = BRAND_PRIMARY
            run.font.size = Pt(15)
            run.bold = True
        elif level == 3:
            run.font.color.rgb = BRAND_DARK
            run.font.size = Pt(12)
            run.bold = True
    return h

def add_callout(doc, text, title="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "EFF6FF")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.color.rgb = BRAND_DARK
    r_text = p.add_run(text)
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 58, 138)
    doc.add_paragraph()

def build_standalone_bi_etl_document():
    print("Generating OmniSuite_BI_ETL_Data_Integration_Documentation.docx...")
    doc = docx.Document()

    # Title Page
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(40)
    title_p.paragraph_format.space_after = Pt(8)
    r = title_p.add_run("OmniSuite ERP")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = BRAND_DARK

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(24)
    r_sub = sub_p.add_run("Business Intelligence & ETL Data Integration Platform\nArchitecture, Technical Specifications & User Guide")
    r_sub.font.size = Pt(16)
    r_sub.font.color.rgb = BRAND_PRIMARY

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(36)
    r_meta = meta_p.add_run("Version 5.0.0 | Release: August 2026\nOmniSuite Core Data & BI Engineering Team")
    r_meta.font.size = Pt(10)
    r_meta.font.color.rgb = TEXT_MUTED

    doc.add_page_break()

    # Section 1: Executive Overview
    add_styled_heading(doc, "1. Executive Overview & Data Platform Vision", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The OmniSuite Business Intelligence (BI) module has been upgraded from static operational reporting "
        "into a full-scale enterprise Business Intelligence and Data Analytics Platform with end-to-end ETL "
        "(Extract, Transform, Load) capabilities. The platform establishes a structured, isolated analytical data layer "
        "enabling high-performance aggregation, historical trend analysis, automated exception detection, and OLAP slicing "
        "without imposing query load on transactional ERP databases."
    )

    add_callout(doc, "The platform follows the standard data engineering lifecycle: EXTRACT -> VALIDATE -> TRANSFORM -> QUALITY CHECK -> LOAD -> ANALYZE -> VISUALIZE -> INSIGHT -> DRILL DOWN -> EXPORT.", "CORE PRINCIPLE")

    # Section 2: Architectural Data Flow
    add_styled_heading(doc, "2. System Architecture & ETL Pipeline Lifecycle", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The BI ETL Engine (server/services/bi/etl.engine.js) orchestrates data pipelines across 6 distinct stages:\n"
        "1. EXTRACT: Queries operational ERP modules (Sales, Purchases, Inventory, GL, Production) with incremental watermarking checkpoints, or parses uploaded Excel (.xlsx, .xls) and CSV datasets.\n"
        "2. VALIDATE: Verifies schema integrity, required primary fields, and type matching.\n"
        "3. TRANSFORM: Applies data cleansing, string trimming, type casting (NUMBER, INTEGER, STRING, BOOLEAN), and calculated column formulas (e.g. gross_profit = net_amount - cost_amount).\n"
        "4. QUALITY CHECK: Computes data quality scores (0-100%) and routes malformed records to bi_rejected_records for quarantine inspection.\n"
        "5. LOAD: Upserts cleansed records into optimized analytical fact tables (bi_fact_sales, bi_fact_purchases, bi_fact_inventory, bi_fact_production, bi_fact_finance).\n"
        "6. COMPLETE: Audits execution duration, stage millisecond timings (extractMs, transformMs, loadMs), and updates pipeline status."
    )

    # Table: Analytical Star Schema Fact Tables
    add_styled_heading(doc, "3. Analytical Star Schema Layer", level=1)
    p = doc.add_paragraph()
    p.add_run("The analytical layer decouples reporting queries from operational tables using a conformed Star Schema:")

    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Fact Table Name", "Grain / Level of Detail", "Key Measures (Metrics)", "Dimension Foreign Keys"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_background(cell, "0E3646")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p_cell = cell.paragraphs[0]
        r_head = p_cell.add_run(h)
        r_head.bold = True
        r_head.font.color.rgb = RGBColor(255, 255, 255)
        r_head.font.size = Pt(9.5)

    facts_data = [
        ("bi_fact_sales", "Invoice / POS Line Item", "quantity, gross_amount, net_amount, cost_amount, gross_profit, margin_pct", "date_key, customer_id, product_id, branch_id, salesperson_id"),
        ("bi_fact_purchases", "Purchase Order Line", "quantity, unit_price, tax_amount, total_amount", "date_key, supplier_id, product_id, branch_id"),
        ("bi_fact_inventory", "Item Warehouse Snapshot", "stock_qty, cost_price, selling_price, total_stock_value, reorder_level", "date_key, product_id, warehouse_id, branch_id"),
        ("bi_fact_production", "Work Order / Job Card", "planned_qty, good_qty, scrap_qty, scrap_rate, yield_rate", "date_key, work_order_id, bom_id, product_id, branch_id"),
        ("bi_fact_finance", "General Ledger Entry", "debit_amount, credit_amount, net_amount", "date_key, account_id, branch_id"),
        ("bi_fact_projects", "Project Monthly Milestone", "budget, total_spent, budget_variance, completion_pct", "date_key, project_id, branch_id")
    ]

    for row_idx, row in enumerate(facts_data):
        r_cells = tbl.add_row().cells
        bg = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            set_cell_background(r_cells[c_idx], bg)
            set_cell_margins(r_cells[c_idx], top=80, bottom=80, left=100, right=100)
            p_c = r_cells[c_idx].paragraphs[0]
            r_val = p_c.add_run(val)
            r_val.font.size = Pt(9)
            if c_idx == 0:
                r_val.bold = True
                r_val.font.color.rgb = BRAND_DARK

    doc.add_paragraph()

    # Section 4: Conformed Date Dimension
    add_styled_heading(doc, "4. Conformed Date Dimension (bi_dim_date)", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The bi_dim_date table provides pre-computed calendar and fiscal attributes spanning 1,827 days from 2024 to 2028. "
        "Attributes include date_key (YYYYMMDD integer), full_date, day_of_month, day_name, month_number, month_name, "
        "quarter_number, quarter_name, year_number, and fiscal_year."
    )

    # Section 5: Multidimensional Analysis Engine (OLAP)
    add_styled_heading(doc, "5. Multidimensional OLAP Slicing Engine", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The multidimensional analytics engine (server/services/bi/analytics.engine.js) provides dynamic slicing across:\n"
        "• Measures: Revenue, Cost of Goods, Gross Profit, Margin %, Units Sold, Spend, Stock Value, Scrap Rate.\n"
        "• Dimensions: Time (Month, Quarter, Year), Branch, Customer, Product, Category, Supplier, Account Category.\n"
        "• Comparisons: Previous Period, Year-over-Year, Budget/Target Benchmarks with automatic variance calculation (Delta Amount and Delta Percentage)."
    )

    # Section 6: Automated Business Exceptions
    add_styled_heading(doc, "6. Automated Anomaly Detection & Business Exceptions", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The anomaly engine scans enterprise fact models to proactively flag operational exceptions:\n"
        "1. Revenue Contractions (Month-over-Month decrease >= 10%)\n"
        "2. Negative Gross Profit Transactions (Selling price < unit cost)\n"
        "3. Low-Stock Safety Threshold Violations (Stock quantity <= Reorder point)\n"
        "4. Overdue Accounts Receivable (> 30 days past invoice due date)\n"
        "5. Manufacturing Scrap Rate Spikes (Scrap rate > 5.0% on work orders)\n"
        "Each insight includes severity tagging (CRITICAL, WARNING, POSITIVE), root-cause explanations, actionable recommendations, and a 1-click drill-down payload."
    )

    # Section 7: User Guide & Navigation
    add_styled_heading(doc, "7. User Interface Guide: Navigation & Studios", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The BI interface features floating portal navigation dropdowns across four primary groups:\n"
        "1. Executive Dashboard: High-level KPI cards, revenue vs expense trends, and operational snapshots.\n"
        "2. Domain Analytics: 12 specialized dashboards for Finance, Inventory, Purchase, HR, Production, PM, Maintenance, Fleet, Service, POS, Admin, and Cross-Module.\n"
        "3. BI Tools: KPI Center, Custom Dashboards, Report Center, Data Explorer, AI Insights, Ask Banks AI, Alerts, and Settings.\n"
        "4. Data & ETL Integration: Data Sources, Analytical Datasets, Data Preparation Studio, Star Schema Models, ETL Pipelines Manager, Data Quality & Quarantine, Multidimensional Slicing, Automated Exceptions, and Custom Dashboard Builder."
    )

    doc.save("OmniSuite_BI_ETL_Data_Integration_Documentation.docx")
    print("OmniSuite_BI_ETL_Data_Integration_Documentation.docx created successfully.")

def update_user_manual():
    path = "OmniSuite_ERP_User_Manual.docx"
    if not os.path.exists(path):
        print(f"File {path} not found.")
        return
    print(f"Updating {path}...")
    doc = docx.Document(path)
    add_styled_heading(doc, "Section 14: Business Intelligence, ETL & Data Integration User Manual", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The OmniSuite BI Module provides a unified data platform with full ETL integration. Authorized users can "
        "manage data ingestion, build visual transformation recipes, monitor automated ETL pipelines, review quarantined records, "
        "slice multidimensional metrics, and investigate automated business anomalies."
    )

    add_styled_heading(doc, "14.1 Managing Data Sources & File Ingestion", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Navigate to Data & ETL Integration -> Data Sources & Ingestion to view operational ERP sources and external datasets. "
        "Click 'Import File Source' to upload Excel (.xlsx, .xls) or CSV spreadsheets. Use 'Test Connection' to verify source health "
        "and 'Sync Now' to trigger on-demand data extraction."
    )

    add_styled_heading(doc, "14.2 Data Preparation & Transformation Studio", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Navigate to Data Preparation Studio to design custom transformation recipes. Users can add calculated columns with formulas "
        "(e.g., net_amount - cost_amount), clean whitespace, apply conditional filters, and cast data types with live before-and-after previews."
    )

    add_styled_heading(doc, "14.3 Monitoring ETL Pipelines & Execution History", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Navigate to ETL Pipelines Manager to view pipeline status, schedule crons, and the 6-stage execution stepper "
        "(Extract -> Validate -> Transform -> Quality Check -> Load -> Complete). Click 'View History' to inspect millisecond stage timings and logs."
    )

    add_styled_heading(doc, "14.4 Data Quality Scoring & Quarantined Records", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Navigate to Data Quality & Quarantine to inspect enterprise data validation scores (0-100%). Records failing validation rules "
        "are quarantined in bi_rejected_records. Click 'Inspect Payload' to view the raw record and rejection reasons."
    )

    add_styled_heading(doc, "14.5 Multidimensional Slicing & Period Variance", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Navigate to Multidimensional Slicing to select measures (Revenue, Cost, Profit Margin, Quantity) and slice across dimensions "
        "(Month, Quarter, Branch, Customer, Product). The system calculates growth variance percentages and exports multi-sheet Excel reports."
    )

    add_styled_heading(doc, "14.6 Automated Business Exceptions & Anomaly Alerts", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Navigate to Automated Exceptions to view autonomous alerts for revenue drops, negative margins, low-stock reorder emergencies, "
        "and overdue receivables. Click 'Investigate' to open interactive drill-down modals directly into the root-cause data."
    )

    doc.save(path)
    print(f"{path} updated successfully.")

def update_technical_architectures():
    paths = [
        "OmniSuite_ERP_Technical_Documentation_and_System_Architecture.docx",
        "OmniSuite_ERP_Technical_System_Documentation.docx"
    ]
    for path in paths:
        if not os.path.exists(path):
            continue
        print(f"Updating {path}...")
        doc = docx.Document(path)
        add_styled_heading(doc, "BI & ETL Data Integration System Architecture", level=1)
        
        p = doc.add_paragraph()
        p.add_run(
            "The OmniSuite BI Analytical Architecture is designed for high-throughput analytical query processing while maintaining strict "
            "OLTP database isolation. It comprises a decoupled Star Schema, an autonomous node-cron ETL engine, and an in-memory multidimensional OLAP slicer."
        )

        add_styled_heading(doc, "Analytical Database Schema & Tables", level=2)
        p = doc.add_paragraph()
        p.add_run(
            "• bi_data_sources: Metadata, connection credentials, source types (ERP, EXCEL, CSV, REST), and incremental watermarks.\n"
            "• bi_datasets: Catalog of analytical datasets with storage types (FACT_TABLE, VIRTUAL) and lineage metadata.\n"
            "• bi_etl_pipelines: Pipeline execution parameters, source-to-target dataset mappings, schedule crons, and transformation steps.\n"
            "• bi_etl_pipeline_runs: Execution history audit trail, duration, stage millisecond timings (stage_timings JSON), and quality scores.\n"
            "• bi_etl_run_logs: Granular stage execution logs (INFO, WARN, ERROR).\n"
            "• bi_data_quality_logs: Validation rules, evaluation checks, and violation counters.\n"
            "• bi_rejected_records: Quarantined records with raw payload JSON and rejection failure reasons.\n"
            "• bi_dim_date: 1,827 days conformed date dimension (2024-2028).\n"
            "• bi_fact_sales, bi_fact_purchases, bi_fact_inventory, bi_fact_production, bi_fact_finance, bi_fact_projects: Analytical fact tables.\n"
            "• bi_custom_dashboards, bi_dashboard_widgets: Custom modular layout storage."
        )

        add_styled_heading(doc, "Incremental Watermarking & Extraction Logic", level=2)
        p = doc.add_paragraph()
        p.add_run(
            "The ETL engine supports high-watermark incremental extraction. When checkpoint_field (e.g. updated_at) is specified, "
            "the extractor queries WHERE updated_at > :lastCheckpoint. Upon successful load, the highest timestamp is persisted as last_checkpoint_value, "
            "preventing redundant transactional extraction."
        )

        doc.save(path)
        print(f"{path} updated successfully.")

def update_api_documentation_docx():
    path = "OmniSuite_ERP_API_Documentation.docx"
    if not os.path.exists(path):
        return
    print(f"Updating {path}...")
    doc = docx.Document(path)
    add_styled_heading(doc, "Business Intelligence & ETL REST API Endpoints", level=1)

    endpoints = [
        ("GET /api/bi/data-sources", "List configured ERP and external file data sources with connection status and record counts."),
        ("POST /api/bi/data-sources", "Create a new operational or external data source configuration."),
        ("POST /api/bi/data-sources/:id/test", "Test connectivity and schema access for a data source."),
        ("POST /api/bi/data-sources/:id/sync", "Trigger an immediate on-demand synchronization of a data source."),
        ("POST /api/bi/upload-source-file", "Upload Excel (.xlsx, .xls) or CSV files for automated ingestion and dataset registration."),
        ("GET /api/bi/datasets", "Retrieve analytical dataset catalog, storage models, row counts, and data freshness."),
        ("GET /api/bi/datasets/:id/preview", "Sample latest tabular records from an analytical dataset."),
        ("POST /api/bi/data-preparation/preview-transform", "Execute transformation steps on sample data with live quality checks."),
        ("GET /api/bi/etl-pipelines", "List ETL pipeline definitions, target fact tables, and schedule crons."),
        ("POST /api/bi/etl-pipelines/:id/run", "Execute full 6-stage ETL pipeline run with stage logging and fact table loading."),
        ("GET /api/bi/etl-pipelines/:id/runs", "Get pipeline execution run history and quality scores."),
        ("GET /api/bi/etl-pipeline-runs/:runId/logs", "Retrieve granular stage execution logs for a specific run."),
        ("GET /api/bi/data-quality/summary", "Get overall data quality score %, rule violations, and quarantined rejected records."),
        ("POST /api/bi/multidimensional-analysis", "Execute multidimensional OLAP aggregation with measures, dimensions, and period comparisons."),
        ("GET /api/bi/insights/automated", "Retrieve automated anomaly detections, variance alerts, and drill-down payloads."),
        ("GET /api/bi/custom-dashboards", "List custom BI dashboards and widget configurations."),
        ("POST /api/bi/custom-dashboards", "Create and save a new custom BI dashboard with modular widgets."),
        ("POST /api/bi/export-custom", "Generate multi-sheet formatted Excel or CSV export files from analytical queries.")
    ]

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Endpoint & Method", "Description & Capabilities"]):
        cell = tbl.cell(0, i)
        set_cell_background(cell, "0E3646")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)

    for row_idx, (ep, desc) in enumerate(endpoints):
        r_cells = tbl.add_row().cells
        bg = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate([ep, desc]):
            set_cell_background(r_cells[c_idx], bg)
            set_cell_margins(r_cells[c_idx], top=80, bottom=80, left=100, right=100)
            p_c = r_cells[c_idx].paragraphs[0]
            r_val = p_c.add_run(val)
            r_val.font.size = Pt(9)
            if c_idx == 0:
                r_val.bold = True
                r_val.font.color.rgb = BRAND_DARK

    doc.save(path)
    print(f"{path} updated successfully.")

def update_openapi_files():
    print("Updating OpenAPI specifications (JSON & YAML)...")
    json_path = "OmniSuite_ERP_OpenAPI.json"
    yaml_path = "OmniSuite_ERP_OpenAPI.yaml"

    with open(json_path, "r", encoding="utf-8") as f:
        spec = json.load(f)

    # Ensure tag exists
    tag_names = [t["name"] for t in spec.get("tags", [])]
    if "Business Intelligence & ETL" not in tag_names:
        spec["tags"].append({
            "name": "Business Intelligence & ETL",
            "description": "Data ingestion, ETL pipelines, star schema fact tables, data quality quarantine, OLAP analysis, and automated insights"
        })

    # Add paths
    bi_paths = {
        "/bi/data-sources": {
            "get": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "List all configured ERP and external data sources",
                "responses": { "200": { "description": "List of data sources" } }
            },
            "post": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Create a new data source configuration",
                "responses": { "201": { "description": "Data source created" } }
            }
        },
        "/bi/data-sources/{id}/test": {
            "post": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Test connection and connectivity health for a data source",
                "parameters": [{ "name": "id", "in": "path", "required": True, "schema": { "type": "integer" } }],
                "responses": { "200": { "description": "Connection test result" } }
            }
        },
        "/bi/data-sources/{id}/sync": {
            "post": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Trigger immediate synchronization of a data source",
                "parameters": [{ "name": "id", "in": "path", "required": True, "schema": { "type": "integer" } }],
                "responses": { "200": { "description": "Sync completed" } }
            }
        },
        "/bi/upload-source-file": {
            "post": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Upload Excel or CSV file for dataset ingestion",
                "requestBody": {
                    "required": True,
                    "content": {
                        "multipart/form-data": {
                            "schema": {
                                "type": "object",
                                "properties": {
                                    "file": { "type": "string", "format": "binary" }
                                }
                            }
                        }
                    }
                },
                "responses": { "200": { "description": "File parsed and dataset registered" } }
            }
        },
        "/bi/datasets": {
            "get": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Get catalog of registered analytical datasets",
                "responses": { "200": { "description": "Catalog of datasets" } }
            }
        },
        "/bi/datasets/{id}/preview": {
            "get": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Sample latest tabular records from an analytical dataset",
                "parameters": [{ "name": "id", "in": "path", "required": True, "schema": { "type": "integer" } }],
                "responses": { "200": { "description": "Dataset sample preview rows" } }
            }
        },
        "/bi/data-preparation/preview-transform": {
            "post": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Preview data transformation recipe with live quality checks",
                "responses": { "200": { "description": "Transformed records and quality score" } }
            }
        },
        "/bi/etl-pipelines": {
            "get": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "List all ETL pipelines and schedules",
                "responses": { "200": { "description": "List of pipelines" } }
            }
        },
        "/bi/etl-pipelines/{id}/run": {
            "post": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Execute an ETL pipeline run",
                "parameters": [{ "name": "id", "in": "path", "required": True, "schema": { "type": "integer" } }],
                "responses": { "200": { "description": "Pipeline run execution results" } }
            }
        },
        "/bi/data-quality/summary": {
            "get": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Get data quality summary score and quarantined records",
                "responses": { "200": { "description": "Data quality score and quarantined rejected records" } }
            }
        },
        "/bi/multidimensional-analysis": {
            "post": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Execute multidimensional OLAP aggregation query",
                "responses": { "200": { "description": "Aggregated multidimensional items and summary" } }
            }
        },
        "/bi/insights/automated": {
            "get": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Get automated anomaly detections and operational exceptions",
                "responses": { "200": { "description": "List of active business insights" } }
            }
        },
        "/bi/custom-dashboards": {
            "get": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "List custom BI dashboards",
                "responses": { "200": { "description": "List of custom dashboards" } }
            },
            "post": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Create a new custom dashboard with modular widgets",
                "responses": { "201": { "description": "Custom dashboard created" } }
            }
        },
        "/bi/export-custom": {
            "post": {
                "tags": ["Business Intelligence & ETL"],
                "summary": "Export analytical data to Excel (multi-sheet) or CSV format",
                "responses": { "200": { "description": "File stream" } }
            }
        }
    }

    spec["paths"].update(bi_paths)

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(spec, f, indent=2)

    with open(yaml_path, "w", encoding="utf-8") as f:
        yaml.dump(spec, f, sort_keys=False)

    print("OpenAPI JSON & YAML updated successfully.")

if __name__ == "__main__":
    build_standalone_bi_etl_document()
    update_user_manual()
    update_technical_architectures()
    update_api_documentation_docx()
    update_openapi_files()
    print("ALL DOCUMENTATION UPDATED SUCCESSFULLY!")
